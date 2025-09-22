import os

from docx import Document


def generate_report(target, outdir, parsed):

    doc = Document()
    doc.add_heading(f"Pentest Report - {target}", 0)
    doc.add_paragraph(f"Target: {target}")

    doc.add_heading("Nmap Summary", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = 'Port'
    hdr[1].text = 'Protocol'
    hdr[2].text = 'Service'
    hdr[3].text = 'Version'
    for p in parsed.get('ports', []):
        row = table.add_row().cells
        row[0].text = str(p['port'])
        row[1].text = p['protocol']
        row[2].text = p['service']
        row[3].text = p['version']

    out_path = os.path.join(outdir, f"report_{target}.docx")
    doc.save(out_path)
    return out_path
